"""DOCX and PDF generation from a Markdown-ish structured text."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from services.formatting_units import DXA_PER_INCH, HALF_POINTS_PER_POINT, TWIPS_PER_POINT
from services.rule_resolver import SYSTEM_DEFAULTS


def _dxa_to_inches(value: Any) -> float:
    try:
        return float(value) / DXA_PER_INCH
    except (TypeError, ValueError):
        return 0.0


def _halfpt_to_pt(value: Any) -> float:
    try:
        return float(value) / HALF_POINTS_PER_POINT
    except (TypeError, ValueError):
        return 0.0


def _twips_to_pt(value: Any) -> float:
    try:
        return float(value) / TWIPS_PER_POINT
    except (TypeError, ValueError):
        return 0.0


def _docx_alignment(value: Any) -> WD_ALIGN_PARAGRAPH:
    normalized = str(value or "").strip().lower()
    if normalized in {"both", "justified", "justify"}:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if normalized == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if normalized == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _reportlab_alignment(value: Any) -> int:
    normalized = str(value or "").strip().lower()
    if normalized in {"both", "justified", "justify"}:
        return TA_JUSTIFY
    if normalized == "center":
        return TA_CENTER
    if normalized == "right":
        return TA_RIGHT
    return TA_LEFT


def _reportlab_font_name(value: Any) -> str:
    normalized = str(value or "").strip().lower()
    if "times" in normalized:
        return "Times-Roman"
    if "courier" in normalized:
        return "Courier"
    return "Helvetica"


def _cover_page_summary(text: str, limit: int = 160) -> str:
    preview = " ".join(str(text or "").split())
    if len(preview) <= limit:
        return preview
    return preview[: limit - 3].rstrip() + "..."


def _page_number_text(page_number: int, page_format: Any) -> str:
    normalized = str(page_format or "").strip().lower()
    if normalized in {"roman", "lowerroman"}:
        return _to_roman(page_number).lower()
    if normalized == "upperroman":
        return _to_roman(page_number).upper()
    return str(page_number)


def _to_roman(page_number: int) -> str:
    if page_number <= 0:
        return ""

    numerals = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    remaining = page_number
    result: list[str] = []
    for value, numeral in numerals:
        while remaining >= value:
            result.append(numeral)
            remaining -= value
    return "".join(result)


def _add_docx_page_field(paragraph) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    paragraph.add_run("1")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _apply_docx_section_geometry(section, formatting: dict[str, Any]) -> None:
    section.page_width = Inches(_dxa_to_inches(formatting["page_width_dxa"]))
    section.page_height = Inches(_dxa_to_inches(formatting["page_height_dxa"]))
    section.top_margin = Inches(_dxa_to_inches(formatting["margin_top_dxa"]))
    section.bottom_margin = Inches(_dxa_to_inches(formatting["margin_bottom_dxa"]))
    section.left_margin = Inches(_dxa_to_inches(formatting["margin_left_dxa"]))
    section.right_margin = Inches(_dxa_to_inches(formatting["margin_right_dxa"]))
    section.header_distance = Inches(_dxa_to_inches(formatting["margin_header_dxa"]))
    section.footer_distance = Inches(_dxa_to_inches(formatting["margin_footer_dxa"]))


def _document_split_requested(formatting: dict[str, Any]) -> bool:
    return bool(
        formatting.get("has_cover_section")
        or formatting.get("has_prelim_section")
        or formatting.get("has_cover_image")
    )


def _resolve_signature_block(compiled_rules: dict[str, Any] | None) -> dict[str, Any]:
    compiled = compiled_rules or {}
    resolved = compiled.get("resolved_rules") if isinstance(compiled.get("resolved_rules"), dict) else {}
    signature = compiled.get("signature_block") if isinstance(compiled.get("signature_block"), dict) else {}

    def pick(*values: Any, default: str = "") -> str:
        for value in values:
            if value is None:
                continue
            text = str(value).strip()
            if text:
                return text
        return default

    return {
        "enabled": bool(signature.get("enabled") or compiled.get("has_signature_block") or resolved.get("has_signature_block")),
        "left_label": pick(signature.get("left_label"), compiled.get("signature_block_left_label"), resolved.get("signature_block_left_label"), default="Prepared By"),
        "right_label": pick(signature.get("right_label"), compiled.get("signature_block_right_label"), resolved.get("signature_block_right_label"), default="Approved By"),
        "left_name": pick(signature.get("left_name"), compiled.get("signature_block_left_name"), resolved.get("signature_block_left_name")),
        "right_name": pick(signature.get("right_name"), compiled.get("signature_block_right_name"), resolved.get("signature_block_right_name")),
    }


def _signature_column_width_dxa(formatting: dict[str, Any]) -> int:
    page_width_dxa = int(formatting.get("page_width_dxa") or SYSTEM_DEFAULTS["page_width_dxa"])
    margin_left_dxa = int(formatting.get("margin_left_dxa") or SYSTEM_DEFAULTS["margin_left_dxa"])
    margin_right_dxa = int(formatting.get("margin_right_dxa") or SYSTEM_DEFAULTS["margin_right_dxa"])
    return max(0, (page_width_dxa - margin_left_dxa - margin_right_dxa) // 2)


def _append_docx_signature_block(doc: Document, formatting: dict[str, Any], signature_block: dict[str, Any]) -> None:
    if not signature_block.get("enabled"):
        return

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(float(formatting["signature_block_space_before_pt"]))
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    column_width = _signature_column_width_dxa(formatting)
    blocks = [
        (table.cell(0, 0), signature_block.get("left_label") or "Prepared By", signature_block.get("left_name") or ""),
        (table.cell(0, 1), signature_block.get("right_label") or "Approved By", signature_block.get("right_name") or ""),
    ]

    for cell, label, name in blocks:
        cell.width = Inches(_dxa_to_inches(column_width))
        first = cell.paragraphs[0]
        first.alignment = _docx_alignment(formatting["signature_alignment"])
        label_run = first.add_run(str(label))
        label_run.font.name = str(formatting["body_font"])
        label_run.bold = True
        label_run.font.size = Pt(float(formatting["signature_label_size_pt"]))

        line = cell.add_paragraph()
        line.alignment = _docx_alignment(formatting["signature_alignment"])
        line_run = line.add_run("____________________")
        line_run.font.name = str(formatting["body_font"])
        line_run.font.size = Pt(float(formatting["signature_line_size_pt"]))

        if name:
            name_para = cell.add_paragraph()
            name_para.alignment = _docx_alignment(formatting["signature_alignment"])
            name_run = name_para.add_run(str(name))
            name_run.font.name = str(formatting["body_font"])
            name_run.font.size = Pt(float(formatting["signature_name_size_pt"]))
            name_run.italic = bool(formatting["signature_name_italic"])


def _append_pdf_signature_block(flow: list[Any], styles: dict[str, ParagraphStyle], formatting: dict[str, Any], signature_block: dict[str, Any]) -> None:
    if not signature_block.get("enabled"):
        return

    centered = ParagraphStyle(
        "SignatureCenter",
        parent=styles["p"],
        alignment=_reportlab_alignment(formatting["signature_alignment"]),
        spaceBefore=float(formatting["signature_block_space_before_pt"]),
    )
    data = [
        [Paragraph(_escape(str(signature_block.get("left_label") or "Prepared By")), centered), Paragraph(_escape(str(signature_block.get("right_label") or "Approved By")), centered)],
        [Paragraph(_escape("____________________"), centered), Paragraph(_escape("____________________"), centered)],
    ]
    left_name = str(signature_block.get("left_name") or "").strip()
    right_name = str(signature_block.get("right_name") or "").strip()
    if left_name or right_name:
        data.append(
            [
                Paragraph(_escape(left_name), centered),
                Paragraph(_escape(right_name), centered),
            ]
        )

    column_width = _dxa_to_inches(_signature_column_width_dxa(formatting)) * inch
    table = Table(data, colWidths=[column_width, column_width])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    flow.append(Spacer(1, float(formatting["signature_block_space_before_pt"])))
    flow.append(table)


def _configure_docx_page_numbers(
    section,
    formatting: dict[str, Any],
    page_format_override: Any | None = None,
    section_restart: bool | None = None,
) -> None:
    if not formatting.get("has_page_numbers"):
        return

    page_format = str(
        page_format_override
        or formatting.get("page_number_format")
        or formatting.get("body_page_format")
        or formatting.get("prelim_page_format")
        or SYSTEM_DEFAULTS.get("body_page_format")
    ).strip().lower()
    format_map = {
        "roman": "lowerRoman",
        "lowerroman": "lowerRoman",
        "upperroman": "upperRoman",
        "decimal": "decimal",
    }
    resolved_format = format_map.get(page_format, "decimal")

    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), resolved_format)
    if section_restart is None:
        section_restart = bool(formatting.get("page_number_section_restart"))
    if section_restart:
        pg_num_type.set(qn("w:start"), "1")

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    paragraph.alignment = _docx_alignment(formatting.get("page_number_alignment"))
    paragraph.add_run("Page ")
    _add_docx_page_field(paragraph)


def _resolve_docx_formatting_rules(compiled_rules: dict[str, Any] | None) -> dict[str, Any]:
    compiled = compiled_rules or {}
    typography = compiled.get("typography") or {}
    layout = compiled.get("layout") or {}
    resolved = compiled.get("resolved_rules") if isinstance(compiled.get("resolved_rules"), dict) else {}

    # Allow passing flattened resolved rules directly as compiled_rules.
    flat_candidate = {
        key: compiled.get(key)
        for key in (
            "page_width_dxa",
            "page_height_dxa",
            "margin_top_dxa",
            "margin_bottom_dxa",
            "margin_left_dxa",
            "margin_right_dxa",
            "margin_header_dxa",
            "margin_footer_dxa",
            "body_font",
            "body_size_halfpt",
            "body_alignment",
            "body_line_spacing_val",
            "body_space_after",
            "headings",
        )
    }
    # Adding page numbering and paper-size fields
    flat_candidate.update({
        "has_page_numbers": compiled.get("has_page_numbers"),
        "page_number_format": compiled.get("page_number_format"),
        "body_page_format": compiled.get("body_page_format"),
        "prelim_page_format": compiled.get("prelim_page_format"),
        "page_number_alignment": compiled.get("page_number_alignment"),
        "page_number_section_restart": compiled.get("page_number_section_restart"),
    })
    has_flat = any(v is not None for v in flat_candidate.values())
    if has_flat and not resolved:
        resolved = flat_candidate

    heading_defaults = {
        "1": SYSTEM_DEFAULTS["headings"]["1"],
        "2": SYSTEM_DEFAULTS["headings"]["2"],
        "3": SYSTEM_DEFAULTS["headings"]["3"],
    }
    headings = resolved.get("headings") if isinstance(resolved.get("headings"), dict) else {}

    def pick(key: str) -> Any:
        value = resolved.get(key)
        return SYSTEM_DEFAULTS[key] if value is None else value

    spacing_factors = SYSTEM_DEFAULTS.get("body_line_spacing_factors") if isinstance(SYSTEM_DEFAULTS.get("body_line_spacing_factors"), dict) else {}
    line_spacing_mode = str(typography.get("line_spacing") or resolved.get("body_line_spacing_rule") or "").strip().lower()
    body_line_spacing_factor = spacing_factors.get(line_spacing_mode)
    if body_line_spacing_factor is None:
        spacing_base = resolved.get("body_line_spacing_val")
        if spacing_base is None:
            spacing_base = SYSTEM_DEFAULTS["body_line_spacing_val"]
        try:
            body_line_spacing_factor = float(spacing_base) / 240.0
        except (TypeError, ValueError):
            body_line_spacing_factor = float(spacing_factors.get("single", 1.0))

    return {
        "page_width_dxa": resolved.get("page_width_dxa") or SYSTEM_DEFAULTS["page_width_dxa"],
        "page_height_dxa": resolved.get("page_height_dxa") or SYSTEM_DEFAULTS["page_height_dxa"],
        "margin_top_dxa": resolved.get("margin_top_dxa") or SYSTEM_DEFAULTS["margin_top_dxa"],
        "margin_bottom_dxa": resolved.get("margin_bottom_dxa") or SYSTEM_DEFAULTS["margin_bottom_dxa"],
        "margin_left_dxa": resolved.get("margin_left_dxa") or SYSTEM_DEFAULTS["margin_left_dxa"],
        "margin_right_dxa": resolved.get("margin_right_dxa") or SYSTEM_DEFAULTS["margin_right_dxa"],
        "margin_header_dxa": resolved.get("margin_header_dxa") or SYSTEM_DEFAULTS["margin_header_dxa"],
        "margin_footer_dxa": resolved.get("margin_footer_dxa") or SYSTEM_DEFAULTS["margin_footer_dxa"],
        "body_font": resolved.get("body_font") or typography.get("font_family") or SYSTEM_DEFAULTS["body_font"],
        "body_size_halfpt": resolved.get("body_size_halfpt") or SYSTEM_DEFAULTS["body_size_halfpt"],
        "body_alignment": resolved.get("body_alignment") or typography.get("alignment") or SYSTEM_DEFAULTS["body_alignment"],
        "body_line_spacing_val": resolved.get("body_line_spacing_val") or SYSTEM_DEFAULTS["body_line_spacing_val"],
        "body_line_spacing_factor": body_line_spacing_factor,
        "body_space_before": pick("body_space_before"),
        "body_space_after": resolved.get("body_space_after") or SYSTEM_DEFAULTS["body_space_after"],
        "heading_numbering": bool(layout.get("heading_numbering")),
        "has_page_numbers": bool(resolved.get("has_page_numbers")),
        "page_number_format": resolved.get("page_number_format") or resolved.get("body_page_format") or SYSTEM_DEFAULTS["body_page_format"],
        "body_page_format": resolved.get("body_page_format") or SYSTEM_DEFAULTS["body_page_format"],
        "prelim_page_format": resolved.get("prelim_page_format") or SYSTEM_DEFAULTS["prelim_page_format"],
        "page_number_alignment": resolved.get("page_number_alignment") or SYSTEM_DEFAULTS["page_number_alignment"],
        "page_number_section_restart": bool(resolved.get("page_number_section_restart")),
        "has_cover_section": bool(resolved.get("has_cover_section")),
        "has_prelim_section": bool(resolved.get("has_prelim_section")),
        "has_body_section": bool(resolved.get("has_body_section", True)),
        "has_cover_image": bool(resolved.get("has_cover_image")),
        "cover_title_size_pt": pick("cover_title_size_pt"),
        "cover_title_leading_factor": pick("cover_title_leading_factor"),
        "cover_title_alignment": pick("cover_title_alignment"),
        "cover_title_space_after_pt": pick("cover_title_space_after_pt"),
        "cover_title_bold": pick("cover_title_bold"),
        "cover_subtitle_size_pt": pick("cover_subtitle_size_pt"),
        "cover_subtitle_leading_factor": pick("cover_subtitle_leading_factor"),
        "cover_subtitle_alignment": pick("cover_subtitle_alignment"),
        "cover_subtitle_space_after_pt": pick("cover_subtitle_space_after_pt"),
        "cover_subtitle_italic": pick("cover_subtitle_italic"),
        "cover_summary_size_pt": pick("cover_summary_size_pt"),
        "cover_summary_leading_factor": pick("cover_summary_leading_factor"),
        "cover_summary_alignment": pick("cover_summary_alignment"),
        "cover_summary_space_after_pt": pick("cover_summary_space_after_pt"),
        "signature_alignment": pick("signature_alignment"),
        "signature_label_size_pt": pick("signature_label_size_pt"),
        "signature_label_bold": pick("signature_label_bold"),
        "signature_line_size_pt": pick("signature_line_size_pt"),
        "signature_name_size_pt": pick("signature_name_size_pt"),
        "signature_name_italic": pick("signature_name_italic"),
        "signature_block_space_before_pt": pick("signature_block_space_before_pt"),
        "heading_leading_factor": pick("heading_leading_factor"),
        "list_left_indent_pt": pick("list_left_indent_pt"),
        "list_block_spacing_pt": pick("list_block_spacing_pt"),
        "headings": {
            level: {
                "font": (headings.get(level) or {}).get("font", heading_defaults[level]["font"]),
                "size_halfpt": (headings.get(level) or {}).get("size_halfpt", heading_defaults[level]["size_halfpt"]),
                "alignment": (headings.get(level) or {}).get("alignment", heading_defaults[level]["alignment"]),
                "space_before": (headings.get(level) or {}).get("space_before", heading_defaults[level]["space_before"]),
                "space_after": (headings.get(level) or {}).get("space_after", heading_defaults[level]["space_after"]),
                "bold": (headings.get(level) or {}).get("bold", heading_defaults[level]["bold"]),
                "italic": (headings.get(level) or {}).get("italic", heading_defaults[level]["italic"]),
                "underline": (headings.get(level) or {}).get("underline", heading_defaults[level]["underline"]),
                "caps": (headings.get(level) or {}).get("caps", heading_defaults[level]["caps"]),
                "small_caps": (headings.get(level) or {}).get("small_caps", heading_defaults[level]["small_caps"]),
                "numbering": (headings.get(level) or {}).get("numbering", heading_defaults[level]["numbering"]),
            }
            for level in heading_defaults
        },
    }


def _parse_blocks(text: str) -> list[tuple[str, str]]:
    """Return list of (kind, text) where kind ∈ {h1, h2, h3, li, oli, p}."""
    blocks: list[tuple[str, str]] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            blocks.append(("blank", ""))
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif re.match(r"^[-*•]\s+", line):
            blocks.append(("li", re.sub(r"^[-*•]\s+", "", line).strip()))
        elif re.match(r"^\d+[.)]\s+", line):
            blocks.append(("oli", re.sub(r"^\d+[.)]\s+", "", line).strip()))
        else:
            blocks.append(("p", line.strip()))
    # collapse paragraph runs
    merged: list[tuple[str, str]] = []
    buf: list[str] = []
    for kind, val in blocks:
        if kind == "p":
            buf.append(val)
        else:
            if buf:
                merged.append(("p", " ".join(buf)))
                buf = []
            if kind != "blank":
                merged.append((kind, val))
    if buf:
        merged.append(("p", " ".join(buf)))
    return merged


def _extract_layout_hints(layout_plan: dict[str, Any] | None, compiled_rules: dict[str, Any] | None) -> dict[str, Any]:
    compiled = compiled_rules or {}
    typography = compiled.get("typography") or {}
    layout = compiled.get("layout") or {}
    plan = layout_plan or {}
    hard_constraints = plan.get("hardConstraintsApplied") or []

    line_spacing = typography.get("line_spacing") or "single"
    alignment = typography.get("alignment") or "justified"
    if any(str(item).startswith("line_spacing:") for item in hard_constraints):
        for item in hard_constraints:
            if str(item).startswith("line_spacing:"):
                line_spacing = str(item).split(":", 1)[1] or line_spacing
                break

    return {
        "line_spacing": line_spacing,
        "alignment": alignment,
        "heading_numbering": bool(layout.get("heading_numbering")),
        "max_heading_depth": int(layout.get("max_heading_depth") or 3),
        "allow_page_breaks": bool(plan.get("placements")),
        "placements": plan.get("placements") or [],
    }


class _SectionNumberer:
    def __init__(self) -> None:
        self._counts = [0, 0, 0]

    def format(self, level: int, title: str) -> str:
        clean = title.strip()
        if re.match(r"^\d+(?:\.\d+)*\s+", clean):
            return clean
        idx = max(1, min(level, 3)) - 1
        self._counts[idx] += 1
        for tail in range(idx + 1, len(self._counts)):
            self._counts[tail] = 0
        parts = [str(n) for n in self._counts[: idx + 1] if n > 0]
        return f"{'.'.join(parts)} {clean}" if parts else clean


# ---------------------------- DOCX -------------------------------------------
def build_docx(
    title: str,
    rules: str,
    structured_text: str,
    out_path: Path,
    layout_plan: dict[str, Any] | None = None,
    compiled_rules: dict[str, Any] | None = None,
) -> None:
    hints = _extract_layout_hints(layout_plan, compiled_rules)
    formatting = _resolve_docx_formatting_rules(compiled_rules)
    doc = Document()
    for section in doc.sections:
        _apply_docx_section_geometry(section, formatting)

    style = doc.styles["Normal"]
    style.font.name = str(formatting["body_font"])
    style.font.size = Pt(_halfpt_to_pt(formatting["body_size_halfpt"]))
    style.paragraph_format.space_before = Pt(_twips_to_pt(formatting["body_space_before"]))
    style.paragraph_format.space_after = Pt(_twips_to_pt(formatting["body_space_after"]))
    style.paragraph_format.line_spacing = float(formatting["body_line_spacing_factor"])
    style.paragraph_format.alignment = _docx_alignment(formatting.get("body_alignment"))

    split_requested = _document_split_requested(formatting)
    prelim_page_format = formatting.get("prelim_page_format") or formatting.get("page_number_format") or formatting.get("body_page_format") or "decimal"
    body_page_format = formatting.get("body_page_format") or formatting.get("page_number_format") or "decimal"

    if split_requested:
        cover = doc.add_paragraph()
        cover.alignment = _docx_alignment(formatting["cover_title_alignment"])
        cover.paragraph_format.space_after = Pt(float(formatting["cover_title_space_after_pt"]))
        cover_run = cover.add_run(title)
        cover_run.font.name = str(formatting["headings"]["1"]["font"])
        cover_run.bold = True
        cover_run.font.size = Pt(float(formatting["cover_title_size_pt"]))
        cover_run.bold = bool(formatting["cover_title_bold"])

        sub = doc.add_paragraph()
        sub.alignment = _docx_alignment(formatting["cover_subtitle_alignment"])
        sub.paragraph_format.space_after = Pt(float(formatting["cover_subtitle_space_after_pt"]))
        sub_run = sub.add_run("Generated by DocuForge AI")
        sub_run.font.name = str(formatting["body_font"])
        sub_run.italic = True
        sub_run.italic = bool(formatting["cover_subtitle_italic"])
        sub_run.font.size = Pt(float(formatting["cover_subtitle_size_pt"]))

        preview = _cover_page_summary(rules)
        if not formatting["has_prelim_section"] and preview:
            summary = doc.add_paragraph()
            summary.alignment = _docx_alignment(formatting["cover_summary_alignment"])
            summary.paragraph_format.space_after = Pt(float(formatting["cover_summary_space_after_pt"]))
            summary_run = summary.add_run(preview)
            summary_run.font.name = str(formatting["body_font"])
            summary_run.font.size = Pt(float(formatting["cover_summary_size_pt"]))

        if formatting["has_prelim_section"]:
            prelim_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            _apply_docx_section_geometry(prelim_section, formatting)
            prelim_section.header.is_linked_to_previous = False
            prelim_section.footer.is_linked_to_previous = False

            prelim_title = doc.add_paragraph()
            prelim_title.alignment = _docx_alignment(formatting["cover_summary_alignment"])
            prelim_title.paragraph_format.space_after = Pt(float(formatting["cover_summary_space_after_pt"]))
            prelim_run = prelim_title.add_run("Preliminary")
            prelim_run.font.name = str(formatting["headings"]["1"]["font"])
            prelim_run.bold = True
            prelim_run.font.size = Pt(float(formatting["cover_title_size_pt"]))

            if preview:
                summary = doc.add_paragraph()
                summary.alignment = _docx_alignment(formatting["cover_summary_alignment"])
                summary.paragraph_format.space_after = Pt(float(formatting["cover_summary_space_after_pt"]))
                summary_run = summary.add_run(preview)
                summary_run.font.name = str(formatting["body_font"])
                summary_run.font.size = Pt(float(formatting["cover_summary_size_pt"]))

            if formatting["has_page_numbers"]:
                _configure_docx_page_numbers(prelim_section, formatting, page_format_override=prelim_page_format, section_restart=True)

        body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _apply_docx_section_geometry(body_section, formatting)
        body_section.header.is_linked_to_previous = False
        body_section.footer.is_linked_to_previous = False

        if formatting["has_page_numbers"]:
            # Leave the cover unnumbered.
            _configure_docx_page_numbers(body_section, formatting, page_format_override=body_page_format, section_restart=True)
    else:
        _configure_docx_page_numbers(doc.sections[0], formatting)

        t = doc.add_paragraph()
        t.alignment = _docx_alignment(formatting["cover_title_alignment"])
        t.paragraph_format.space_after = Pt(float(formatting["cover_title_space_after_pt"]))
        run = t.add_run(title)
        run.font.name = str(formatting["headings"]["1"]["font"])
        run.bold = True
        run.font.size = Pt(float(formatting["cover_title_size_pt"]))
        run.bold = bool(formatting["cover_title_bold"])

        sub = doc.add_paragraph()
        sub.alignment = _docx_alignment(formatting["cover_subtitle_alignment"])
        sub.paragraph_format.space_after = Pt(float(formatting["cover_subtitle_space_after_pt"]))
        sub_run = sub.add_run("Generated by DocuForge AI")
        sub_run.font.name = str(formatting["body_font"])
        sub_run.italic = True
        sub_run.italic = bool(formatting["cover_subtitle_italic"])
        sub_run.font.size = Pt(float(formatting["cover_subtitle_size_pt"]))

        doc.add_paragraph()

    blocks = _parse_blocks(structured_text)
    placements = hints["placements"]
    numberer = _SectionNumberer()
    current_page = None

    h1_style = doc.styles["Heading 1"]
    h2_style = doc.styles["Heading 2"]
    h3_style = doc.styles["Heading 3"]

    h1_rules = formatting["headings"]["1"]
    h2_rules = formatting["headings"]["2"]
    h3_rules = formatting["headings"]["3"]

    h1_style.font.name = str(h1_rules["font"])
    h1_style.font.size = Pt(_halfpt_to_pt(h1_rules["size_halfpt"]))
    h1_style.paragraph_format.alignment = _docx_alignment(h1_rules["alignment"])
    h1_style.font.bold = bool(h1_rules.get("bold"))
    h1_style.font.italic = bool(h1_rules.get("italic"))
    h1_style.font.underline = bool(h1_rules.get("underline"))
    h1_style.font.all_caps = bool(h1_rules.get("caps"))
    h1_style.font.small_caps = bool(h1_rules.get("small_caps"))
    h1_style.paragraph_format.space_before = Pt(_twips_to_pt(h1_rules["space_before"]))
    h1_style.paragraph_format.space_after = Pt(_twips_to_pt(h1_rules["space_after"]))

    h2_style.font.name = str(h2_rules["font"])
    h2_style.font.size = Pt(_halfpt_to_pt(h2_rules["size_halfpt"]))
    h2_style.paragraph_format.alignment = _docx_alignment(h2_rules["alignment"])
    h2_style.font.bold = bool(h2_rules.get("bold"))
    h2_style.font.italic = bool(h2_rules.get("italic"))
    h2_style.font.underline = bool(h2_rules.get("underline"))
    h2_style.font.all_caps = bool(h2_rules.get("caps"))
    h2_style.font.small_caps = bool(h2_rules.get("small_caps"))
    h2_style.paragraph_format.space_before = Pt(_twips_to_pt(h2_rules["space_before"]))
    h2_style.paragraph_format.space_after = Pt(_twips_to_pt(h2_rules["space_after"]))

    h3_style.font.name = str(h3_rules["font"])
    h3_style.font.size = Pt(_halfpt_to_pt(h3_rules["size_halfpt"]))
    h3_style.paragraph_format.alignment = _docx_alignment(h3_rules["alignment"])
    h3_style.font.bold = bool(h3_rules.get("bold"))
    h3_style.font.italic = bool(h3_rules.get("italic"))
    h3_style.font.underline = bool(h3_rules.get("underline"))
    h3_style.font.all_caps = bool(h3_rules.get("caps"))
    h3_style.font.small_caps = bool(h3_rules.get("small_caps"))
    h3_style.paragraph_format.space_before = Pt(_twips_to_pt(h3_rules["space_before"]))
    h3_style.paragraph_format.space_after = Pt(_twips_to_pt(h3_rules["space_after"]))

    for index, (kind, val) in enumerate(blocks):
        placement = placements[index] if index < len(placements) else None
        page = int(placement.get("page") or 1) if placement else 1
        if current_page is None:
            current_page = page
        elif hints["allow_page_breaks"] and page > current_page:
            doc.add_page_break()
            current_page = page

        if kind == "h1":
            heading = doc.add_heading(numberer.format(1, val), level=1)
            if hints["heading_numbering"] or formatting["heading_numbering"]:
                heading.style = doc.styles["Heading 1"]
        elif kind == "h2":
            doc.add_heading(numberer.format(2, val), level=2)
        elif kind == "h3":
            doc.add_heading(numberer.format(3, val), level=3)
        elif kind == "li":
            doc.add_paragraph(val, style="List Bullet")
        elif kind == "oli":
            doc.add_paragraph(val, style="List Number")
        else:
            p = doc.add_paragraph(val)
            p.paragraph_format.space_before = Pt(_twips_to_pt(formatting["body_space_before"]))
            p.paragraph_format.space_after = Pt(_twips_to_pt(formatting.get("body_space_after")))
            p.paragraph_format.line_spacing = float(formatting["body_line_spacing_factor"])
            p.paragraph_format.alignment = _docx_alignment(formatting.get("body_alignment"))

    _append_docx_signature_block(doc, formatting, _resolve_signature_block(compiled_rules))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------------------------- PDF --------------------------------------------
def _styles(formatting: dict[str, Any], hints: dict[str, Any] | None = None):
    base = getSampleStyleSheet()
    h1_rules = formatting["headings"]["1"]
    h2_rules = formatting["headings"]["2"]
    h3_rules = formatting["headings"]["3"]
    body_font = _reportlab_font_name(formatting["body_font"])
    h1_font = _reportlab_font_name(h1_rules["font"])
    h2_font = _reportlab_font_name(h2_rules["font"])
    h3_font = _reportlab_font_name(h3_rules["font"])
    body_size = _halfpt_to_pt(formatting["body_size_halfpt"])
    body_leading = body_size * float(formatting["body_line_spacing_factor"])
    h1_size = _halfpt_to_pt(h1_rules["size_halfpt"])
    h2_size = _halfpt_to_pt(h2_rules["size_halfpt"])
    h3_size = _halfpt_to_pt(h3_rules["size_halfpt"])
    title_size = float(formatting["cover_title_size_pt"])
    subtitle_size = float(formatting["cover_subtitle_size_pt"])
    summary_size = float(formatting["cover_summary_size_pt"])
    heading_factor = float(formatting["heading_leading_factor"])
    return {
        "title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName=h1_font,
            fontSize=title_size,
            leading=title_size * float(formatting["cover_title_leading_factor"]),
            spaceAfter=float(formatting["cover_title_space_after_pt"]),
            alignment=_reportlab_alignment(formatting["cover_title_alignment"]),
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            parent=base["Normal"],
            fontName=body_font,
            fontSize=subtitle_size,
            leading=subtitle_size * float(formatting["cover_subtitle_leading_factor"]),
            textColor="#666666",
            alignment=_reportlab_alignment(formatting["cover_subtitle_alignment"]),
            spaceAfter=float(formatting["cover_subtitle_space_after_pt"]),
        ),
        "cover_summary": ParagraphStyle(
            "CoverSummary",
            parent=base["BodyText"],
            fontName=body_font,
            fontSize=summary_size,
            leading=summary_size * float(formatting["cover_summary_leading_factor"]),
            alignment=_reportlab_alignment(formatting["cover_summary_alignment"]),
            spaceAfter=float(formatting["cover_summary_space_after_pt"]),
        ),
        "h1": ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName=h1_font,
            fontSize=h1_size,
            leading=h1_size * heading_factor,
            spaceBefore=_twips_to_pt(h1_rules["space_before"]),
            spaceAfter=_twips_to_pt(h1_rules["space_after"]),
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName=h2_font,
            fontSize=h2_size,
            leading=h2_size * heading_factor,
            spaceBefore=_twips_to_pt(h2_rules["space_before"]),
            spaceAfter=_twips_to_pt(h2_rules["space_after"]),
        ),
        "h3": ParagraphStyle(
            "H3",
            parent=base["Heading3"],
            fontName=h3_font,
            fontSize=h3_size,
            leading=h3_size * heading_factor,
            spaceBefore=_twips_to_pt(h3_rules["space_before"]),
            spaceAfter=_twips_to_pt(h3_rules["space_after"]),
        ),
        "p": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName=body_font,
            fontSize=body_size,
            leading=body_leading,
            spaceBefore=_twips_to_pt(formatting["body_space_before"]),
            spaceAfter=_twips_to_pt(formatting["body_space_after"]),
            alignment=_reportlab_alignment(formatting["body_alignment"]),
        ),
        "li": ParagraphStyle(
            "Li",
            parent=base["BodyText"],
            fontName=body_font,
            fontSize=body_size,
            leading=body_leading,
            spaceBefore=_twips_to_pt(formatting["body_space_before"]),
            spaceAfter=_twips_to_pt(formatting["body_space_after"]),
            leftIndent=float(formatting["list_left_indent_pt"]),
            alignment=_reportlab_alignment(formatting["body_alignment"]),
        ),
        "oli": ParagraphStyle(
            "OLi",
            parent=base["BodyText"],
            fontName=body_font,
            fontSize=body_size,
            leading=body_leading,
            spaceBefore=_twips_to_pt(formatting["body_space_before"]),
            spaceAfter=_twips_to_pt(formatting["body_space_after"]),
            leftIndent=float(formatting["list_left_indent_pt"]),
            alignment=_reportlab_alignment(formatting["body_alignment"]),
        ),
    }


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_pdf(
    title: str,
    rules: str,
    structured_text: str,
    out_path: Path,
    layout_plan: dict[str, Any] | None = None,
    compiled_rules: dict[str, Any] | None = None,
) -> None:
    hints = _extract_layout_hints(layout_plan, compiled_rules)
    formatting = _resolve_docx_formatting_rules(compiled_rules)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    page_width_in = _dxa_to_inches(formatting["page_width_dxa"])
    page_height_in = _dxa_to_inches(formatting["page_height_dxa"])
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=(page_width_in * inch, page_height_in * inch),
        leftMargin=_dxa_to_inches(formatting["margin_left_dxa"]) * inch,
        rightMargin=_dxa_to_inches(formatting["margin_right_dxa"]) * inch,
        topMargin=_dxa_to_inches(formatting["margin_top_dxa"]) * inch,
        bottomMargin=_dxa_to_inches(formatting["margin_bottom_dxa"]) * inch,
        title=title,
        author="DocuForge AI",
    )
    s = _styles(formatting, hints)
    split_requested = _document_split_requested(formatting)
    prelim_page_format = formatting.get("prelim_page_format") or formatting.get("page_number_format") or formatting.get("body_page_format") or "decimal"
    body_page_format = formatting.get("body_page_format") or formatting.get("page_number_format") or "decimal"

    flow = []
    if split_requested:
        flow.extend(
            [
                Paragraph(_escape(title), s["title"]),
                Paragraph("Generated by DocuForge AI", s["subtitle"]),
            ]
        )
        preview = _cover_page_summary(rules)
        if not formatting["has_prelim_section"] and preview:
            flow.append(Paragraph(_escape(preview), s["cover_summary"]))
            flow.append(PageBreak())
        elif not formatting["has_prelim_section"]:
            flow.append(PageBreak())
        elif formatting["has_prelim_section"]:
            flow.append(PageBreak())
            flow.append(Paragraph(_escape("Preliminary"), s["title"]))
            if preview:
                flow.append(Paragraph(_escape(preview), s["cover_summary"]))
            flow.append(PageBreak())
    else:
        flow.extend(
            [
                Paragraph(_escape(title), s["title"]),
                Paragraph("Generated by DocuForge AI", s["subtitle"]),
            ]
        )

    blocks = _parse_blocks(structured_text)
    placements = hints["placements"]
    pending_list: list[ListItem] = []
    pending_olist: list[ListItem] = []
    numberer = _SectionNumberer()
    current_page = None

    def draw_page_number(canvas, doc) -> None:
        if not formatting.get("has_page_numbers"):
            return

        canvas.saveState()
        page_index = canvas.getPageNumber()
        if split_requested and formatting.get("has_cover_section"):
            if formatting.get("has_prelim_section"):
                if page_index == 1:
                    canvas.restoreState()
                    return
                if page_index == 2:
                    page_number = _page_number_text(1, prelim_page_format)
                else:
                    page_number = _page_number_text(page_index - 2, body_page_format)
            else:
                if page_index == 1:
                    canvas.restoreState()
                    return
                page_number = _page_number_text(page_index - 1, body_page_format)
        else:
            page_number = _page_number_text(page_index, formatting.get("page_number_format"))
        label = f"Page {page_number}"
        footer_y = _dxa_to_inches(formatting.get("margin_footer_dxa") or SYSTEM_DEFAULTS["margin_footer_dxa"]) * inch
        alignment = str(formatting.get("page_number_alignment") or SYSTEM_DEFAULTS["page_number_alignment"]).strip().lower()
        if alignment == "left":
            canvas.drawString(doc.leftMargin, footer_y, label)
        elif alignment == "right":
            canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, footer_y, label)
        else:
            canvas.drawCentredString(doc.pagesize[0] / 2.0, footer_y, label)
        canvas.restoreState()

    def flush_list():
        nonlocal pending_list
        if pending_list:
            flow.append(ListFlowable(pending_list, bulletType="bullet", leftIndent=float(formatting["list_left_indent_pt"])))
            flow.append(Spacer(1, float(formatting["list_block_spacing_pt"])))
            pending_list = []

    def flush_olist():
        nonlocal pending_olist
        if pending_olist:
            flow.append(ListFlowable(pending_olist, bulletType="1", leftIndent=float(formatting["list_left_indent_pt"])))
            flow.append(Spacer(1, float(formatting["list_block_spacing_pt"])))
            pending_olist = []

    for index, (kind, val) in enumerate(blocks):
        placement = placements[index] if index < len(placements) else None
        page = int(placement.get("page") or 1) if placement else 1
        if current_page is None:
            current_page = page
        elif hints["allow_page_breaks"] and page > current_page:
            flush_list()
            flush_olist()
            flow.append(PageBreak())
            current_page = page

        if kind == "li":
            flush_olist()
            pending_list.append(ListItem(Paragraph(_escape(val), s["li"])))
            continue
        if kind == "oli":
            bullet = doc.add_paragraph(val, style="List Bullet")
            bullet.paragraph_format.left_indent = Pt(float(formatting["list_left_indent_pt"]))
            bullet.paragraph_format.space_before = Pt(_twips_to_pt(formatting["body_space_before"]))
            bullet.paragraph_format.space_after = Pt(_twips_to_pt(formatting["body_space_after"]))
            pending_olist.append(ListItem(Paragraph(_escape(val), s["oli"])))
            ordered = doc.add_paragraph(val, style="List Number")
            ordered.paragraph_format.left_indent = Pt(float(formatting["list_left_indent_pt"]))
            ordered.paragraph_format.space_before = Pt(_twips_to_pt(formatting["body_space_before"]))
            ordered.paragraph_format.space_after = Pt(_twips_to_pt(formatting["body_space_after"]))
        flush_list()
        flush_olist()
        if kind == "h1":
            flow.append(Paragraph(_escape(numberer.format(1, val)), s["h1"]))
        elif kind == "h2":
            flow.append(Paragraph(_escape(numberer.format(2, val)), s["h2"]))
        elif kind == "h3":
            flow.append(Paragraph(_escape(numberer.format(3, val)), s["h3"]))
        else:
            flow.append(Paragraph(_escape(val), s["p"]))
    flush_list()
    flush_olist()

    _append_pdf_signature_block(flow, s, formatting, _resolve_signature_block(compiled_rules))

    doc.build(flow, onFirstPage=draw_page_number, onLaterPages=draw_page_number)
