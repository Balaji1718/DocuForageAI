from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from pypdf import PdfReader

from doc_generator import build_docx, build_pdf


def test_generator_applies_rule_driven_page_geometry_and_numbering() -> None:
    structured_text = """# Introduction
This is the intro paragraph.

## Body
- first item
- second item

## Conclusion
Final wrap-up paragraph.
"""

    resolved_rules = {
        "page_width_dxa": 11906,  # A4 width
        "page_height_dxa": 16838,  # A4 height
        "margin_top_dxa": 1800,
        "margin_bottom_dxa": 1800,
        "margin_left_dxa": 1440,
        "margin_right_dxa": 1440,
        "margin_header_dxa": 720,
        "margin_footer_dxa": 720,
        "body_font": "Arial",
        "body_size_halfpt": 22,
        "body_alignment": "both",
        "body_line_spacing_val": 360,
        "body_space_after": 120,
        "has_page_numbers": True,
        "page_number_format": "lowerRoman",
        "page_number_alignment": "center",
        "page_number_section_restart": True,
        "headings": {
            "1": {"font": "Calibri", "size_halfpt": 36, "alignment": "left", "space_after": 120},
            "2": {"font": "Calibri", "size_halfpt": 32, "alignment": "left", "space_after": 100},
            "3": {"font": "Calibri", "size_halfpt": 28, "alignment": "left", "space_after": 80},
        },
    }

    with TemporaryDirectory() as td:
        tmp = Path(td)
        docx_path = tmp / "report.docx"
        pdf_path = tmp / "report.pdf"
        compiled_rules = {
            "resolved_rules": resolved_rules,
            "typography": {"line_spacing": "single", "alignment": "justified"},
            "layout": {"heading_numbering": True, "max_heading_depth": 3},
        }

        build_docx(
            title="Render Check",
            rules="Use standard academic structure",
            structured_text=structured_text,
            out_path=docx_path,
            layout_plan={"placements": [{"page": 1}], "totalPages": 1, "hardConstraintsApplied": []},
            compiled_rules=compiled_rules,
        )
        build_pdf(
            title="Render Check",
            rules="Use standard academic structure",
            structured_text=structured_text,
            out_path=pdf_path,
            layout_plan={"placements": [{"page": 1}], "totalPages": 1, "hardConstraintsApplied": []},
            compiled_rules=compiled_rules,
        )

        docx = Document(str(docx_path))
        section = docx.sections[0]
        assert abs(section.page_width.inches - 8.27) < 0.08
        assert abs(section.page_height.inches - 11.69) < 0.08
        assert abs(section.top_margin.inches - 1.25) < 0.08
        assert abs(section.left_margin.inches - 1.0) < 0.08
        assert abs(section.footer_distance.inches - 0.5) < 0.08
        assert "lowerRoman" in section._sectPr.xml
        assert 'w:start="1"' in section._sectPr.xml

        pdf_reader = PdfReader(str(pdf_path))
        first_page = pdf_reader.pages[0]
        width = float(first_page.mediabox.width)
        height = float(first_page.mediabox.height)
        assert abs(width - 595.28) < 2.0
        assert abs(height - 841.89) < 2.0
        pdf_text = first_page.extract_text() or ""
        assert "Page i" in pdf_text


def test_generator_emits_cover_page_and_body_split_when_requested() -> None:
    structured_text = """# Introduction
This is the body paragraph.

## Body
Another body paragraph.
"""

    resolved_rules = {
        "page_width_dxa": 11906,
        "page_height_dxa": 16838,
        "margin_top_dxa": 1800,
        "margin_bottom_dxa": 1800,
        "margin_left_dxa": 1440,
        "margin_right_dxa": 1440,
        "margin_header_dxa": 720,
        "margin_footer_dxa": 720,
        "body_font": "Arial",
        "body_size_halfpt": 22,
        "body_alignment": "both",
        "body_line_spacing_val": 360,
        "body_space_after": 120,
        "has_page_numbers": True,
        "prelim_page_format": "lowerRoman",
        "body_page_format": "decimal",
        "page_number_alignment": "center",
        "page_number_section_restart": True,
        "has_cover_section": True,
        "has_prelim_section": True,
        "has_body_section": True,
        "has_cover_image": False,
        "headings": {
            "1": {"font": "Calibri", "size_halfpt": 36, "alignment": "left", "space_after": 120},
            "2": {"font": "Calibri", "size_halfpt": 32, "alignment": "left", "space_after": 100},
            "3": {"font": "Calibri", "size_halfpt": 28, "alignment": "left", "space_after": 80},
        },
    }

    with TemporaryDirectory() as td:
        tmp = Path(td)
        docx_path = tmp / "split-report.docx"
        pdf_path = tmp / "split-report.pdf"
        compiled_rules = {
            "resolved_rules": resolved_rules,
            "typography": {"line_spacing": "single", "alignment": "justified"},
            "layout": {"heading_numbering": True, "max_heading_depth": 3},
        }

        build_docx(
            title="Split Check",
            rules="Cover page and prelim/body split",
            structured_text=structured_text,
            out_path=docx_path,
            layout_plan={"placements": [{"page": 1}], "totalPages": 1, "hardConstraintsApplied": []},
            compiled_rules=compiled_rules,
        )
        build_pdf(
            title="Split Check",
            rules="Cover page and prelim/body split",
            structured_text=structured_text,
            out_path=pdf_path,
            layout_plan={"placements": [{"page": 1}], "totalPages": 1, "hardConstraintsApplied": []},
            compiled_rules=compiled_rules,
        )

        docx = Document(str(docx_path))
        assert len(docx.sections) >= 3
        cover_section = docx.sections[0]
        prelim_section = docx.sections[1]
        body_section = docx.sections[2]
        assert "pgNumType" not in cover_section._sectPr.xml
        assert "lowerRoman" in prelim_section._sectPr.xml
        assert "decimal" in body_section._sectPr.xml
        assert 'w:start="1"' in body_section._sectPr.xml

        pdf_reader = PdfReader(str(pdf_path))
        assert len(pdf_reader.pages) >= 3
        cover_text = pdf_reader.pages[0].extract_text() or ""
        prelim_text = pdf_reader.pages[1].extract_text() or ""
        body_text = pdf_reader.pages[2].extract_text() or ""
        assert "Split Check" in cover_text
        assert "Page" not in cover_text
        assert "Preliminary" in prelim_text
        assert "Page i" in prelim_text
        assert "Introduction" in body_text
        assert "Page 1" in body_text


def test_generator_appends_two_column_signature_block_when_requested() -> None:
    structured_text = """# Introduction
This body stays short so the signature block is easy to validate.
"""

    resolved_rules = {
        "page_width_dxa": 11906,
        "page_height_dxa": 16838,
        "margin_top_dxa": 1440,
        "margin_bottom_dxa": 1440,
        "margin_left_dxa": 1800,
        "margin_right_dxa": 1800,
        "margin_header_dxa": 720,
        "margin_footer_dxa": 720,
        "body_font": "Arial",
        "body_size_halfpt": 22,
        "body_alignment": "both",
        "body_line_spacing_val": 360,
        "body_space_after": 120,
        "has_page_numbers": False,
        "has_signature_block": True,
        "signature_block_left_label": "Student Signature",
        "signature_block_right_label": "Supervisor Signature",
        "signature_block_left_name": "Balaji",
        "signature_block_right_name": "Dr. Rao",
        "headings": {
            "1": {"font": "Calibri", "size_halfpt": 36, "alignment": "left", "space_after": 120},
            "2": {"font": "Calibri", "size_halfpt": 32, "alignment": "left", "space_after": 100},
            "3": {"font": "Calibri", "size_halfpt": 28, "alignment": "left", "space_after": 80},
        },
    }

    with TemporaryDirectory() as td:
        tmp = Path(td)
        docx_path = tmp / "signature-report.docx"
        pdf_path = tmp / "signature-report.pdf"
        compiled_rules = {
            "resolved_rules": resolved_rules,
            "typography": {"line_spacing": "single", "alignment": "justified"},
            "layout": {"heading_numbering": True, "max_heading_depth": 3},
            "has_signature_block": True,
            "signature_block": {
                "enabled": True,
                "left_label": "Student Signature",
                "right_label": "Supervisor Signature",
                "left_name": "Balaji",
                "right_name": "Dr. Rao",
            },
        }

        build_docx(
            title="Signature Check",
            rules="Signature block validation",
            structured_text=structured_text,
            out_path=docx_path,
            layout_plan={"placements": [{"page": 1}], "totalPages": 1, "hardConstraintsApplied": []},
            compiled_rules=compiled_rules,
        )
        build_pdf(
            title="Signature Check",
            rules="Signature block validation",
            structured_text=structured_text,
            out_path=pdf_path,
            layout_plan={"placements": [{"page": 1}], "totalPages": 1, "hardConstraintsApplied": []},
            compiled_rules=compiled_rules,
        )

        docx = Document(str(docx_path))
        assert len(docx.tables) >= 1
        signature_table = docx.tables[-1]
        assert len(signature_table.columns) == 2
        assert "Student Signature" in signature_table.cell(0, 0).text
        assert "Supervisor Signature" in signature_table.cell(0, 1).text
        assert "tblBorders" not in signature_table._tbl.xml

        pdf_reader = PdfReader(str(pdf_path))
        pdf_text = "\n".join((page.extract_text() or "") for page in pdf_reader.pages)
        assert "Student Signature" in pdf_text
        assert "Supervisor Signature" in pdf_text
        assert "Balaji" in pdf_text
        assert "Dr. Rao" in pdf_text