from __future__ import annotations

import sys
import zipfile
from decimal import Decimal
from pathlib import Path
from types import MappingProxyType
import xml.etree.ElementTree as ET

from docx import Document
import pytest

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from services.generation_v2.docx_fonts import apply_run_font_override
from services.generation_v2.fonts import FONT_CACHE, FontFaceMetrics, build_font_cache
from services.generation_v2.units import pt_to_emu


def test_pt_to_emu_exact_and_integer():
    value = pt_to_emu(12.0)
    assert value == 152400
    assert isinstance(value, int)


def test_font_cache_is_populated_at_startup_and_immutable(monkeypatch):
    fake_font_files = [Path("FakeFont-Regular.ttf")]

    def _fake_iter(_fonts_dir: Path):
        return fake_font_files

    def _fake_loader(_font_path: Path):
        return FontFaceMetrics(
            font_name="FakeFont",
            ascender=1900,
            descender=-500,
            units_per_em=2048,
            cap_height=1450,
            line_height_ratio=Decimal("1.171875"),
            average_advance_width_units=1100,
        )

    monkeypatch.setattr("services.generation_v2.fonts._iter_font_files", _fake_iter)
    monkeypatch.setattr("services.generation_v2.fonts._load_font_face", _fake_loader)

    cache = build_font_cache(Path("fonts"), sizes_pt=[10, 12])

    assert isinstance(cache.by_font_size, MappingProxyType)
    assert len(cache.by_font_size) == 2

    metrics = cache.get("FakeFont", 12)
    assert metrics.font_name == "FakeFont"
    assert metrics.line_height_pt == Decimal("14.0625")

    try:
        cache.by_font_size[("FakeFont", Decimal("12").normalize())] = metrics
        mutated = True
    except TypeError:
        mutated = False
    assert mutated is False


def test_font_override_sets_all_run_font_attributes(tmp_path: Path):
    out_file = tmp_path / "font_override.docx"

    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Deterministic font override test")

    apply_run_font_override(run, "Times New Roman")
    doc.save(out_file)

    with zipfile.ZipFile(out_file, "r") as zf:
        xml_content = zf.read("word/document.xml")

    root = ET.fromstring(xml_content)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    r_fonts = root.find(".//w:rFonts", ns)

    assert r_fonts is not None
    key_ascii = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
    key_hansi = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi"
    key_east_asia = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
    key_cs = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs"

    assert r_fonts.get(key_ascii) == "Times New Roman"
    assert r_fonts.get(key_hansi) == "Times New Roman"
    assert r_fonts.get(key_east_asia) == "Times New Roman"
    assert r_fonts.get(key_cs) == "Times New Roman"


def test_module_level_font_cache_is_immutable():
    with pytest.raises(TypeError):
        FONT_CACHE["test"] = 1
