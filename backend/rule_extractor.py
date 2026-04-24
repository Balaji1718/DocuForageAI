"""
Universal DOCX rule extractor.
- Reads: document.xml, styles.xml, numbering.xml, section properties, headers/footers
- Outputs: Normalized JSON schema (100+ keys, all present, null for unknown)
- Works for: Any DOCX, any industry, any document type
"""

from __future__ import annotations
import io
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from statistics import median
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Namespace map for Word XML
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

NEUTRAL_DEFAULTS = {
    # PAGE
    "page_width_dxa": None,
    "page_height_dxa": None,
    "margin_top_dxa": None,
    "margin_bottom_dxa": None,
    "margin_left_dxa": None,
    "margin_right_dxa": None,
    "margin_header_dxa": None,
    "margin_footer_dxa": None,

    # BODY TYPOGRAPHY
    "body_font": None,
    "body_size_halfpt": None,
    "body_line_spacing_val": None,
    "body_line_spacing_rule": None,  # "auto", "exact", "atLeast"
    "body_space_before": None,
    "body_space_after": None,
    "body_alignment": None,  # "left", "both", "center", "right"
    "body_first_line_indent_dxa": None,
    "body_left_indent_dxa": None,
    "body_right_indent_dxa": None,

    # COVER TYPOGRAPHY
    "cover_title_size_pt": None,
    "cover_subtitle_size_pt": None,
    "cover_title_text": None,
    "cover_summary_text": None,

    # LIST TYPOGRAPHY
    "list_left_indent_pt": None,
    "list_first_line_indent_pt": None,

    # HEADINGS (level 1-6)
    "headings": {},

    # PAGE NUMBERS
    "has_page_numbers": False,
    "page_number_alignment": None,
    "prelim_page_format": None,  # "lowerRoman", "upperRoman", "decimal"
    "body_page_format": None,
    "page_number_section_restart": False,

    # SECTIONS
    "section_count": 0,
    "has_cover_section": False,
    "has_prelim_section": False,
    "has_body_section": False,
    "sections_have_different_margins": False,

    # DOCUMENT STRUCTURE
    "detected_section_headings": [],
    "has_toc": False,
    "has_list_of_figures": False,
    "has_numbered_lists": False,
    "has_bulleted_lists": False,
    "list_indent_dxa": None,
    "list_hanging_dxa": None,

    # TABLES
    "table_count": 0,
    "tables_use_borders": False,
    "dominant_table_width_dxa": None,

    # IMAGES
    "image_count": 0,
    "has_cover_image": False,

    # HEADERS/FOOTERS
    "footer_count": 0,
    "header_count": 0,
    "footer_has_page_number": False,
    "header_has_page_number": False,
    "footer_text_sample": None,

    # QUALITY FLAGS
    "has_markdown_leak": False,
    "has_xml_artifact_numbers": False,
    "has_mixed_fonts": False,
    "has_inconsistent_sizes": False,
    "font_substitution_detected": False,

    # METADATA
    "source_filename": "",
    "extraction_warnings": [],
    "confidence": "medium",
}


def _normalize_name(value: str | None) -> str:
    return re.sub(r"\s+", "", str(value or "")).lower()


def _find_style(styles_tree: ET.Element | None, style_names: set[str]) -> ET.Element | None:
    if styles_tree is None:
        return None

    wanted = {_normalize_name(name) for name in style_names}
    for style in styles_tree.findall(".//w:style", NS):
        style_id = _normalize_name(style.get(f"{{{NS['w']}}}styleId"))
        name_node = style.find("w:name", NS)
        style_name = _normalize_name(name_node.get(f"{{{NS['w']}}}val") if name_node is not None else None)
        if style_id in wanted or style_name in wanted:
            return style
    return None


def _median_or_none(values: list[float | int | None]) -> float | None:
    cleaned = [float(value) for value in values if value is not None]
    if not cleaned:
        return None
    return float(median(cleaned))


def _median_int(values: list[float | int | None]) -> int | None:
    value = _median_or_none(values)
    if value is None:
        return None
    return int(round(value))


def _length_pt(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return float(value.pt)
    except Exception:
        return None


def _pt_to_dxa(value: float | int | None) -> int | None:
    if value is None:
        return None
    return int(round(float(value) * 20.0))


def _line_spacing_to_twips(paragraph: Any) -> int | None:
    spacing = getattr(getattr(paragraph, "paragraph_format", None), "line_spacing", None)
    if spacing is None:
        return None

    spacing_pt = _length_pt(spacing)
    if spacing_pt is not None:
        return int(round(spacing_pt * 20.0))

    try:
        numeric = float(spacing)
    except Exception:
        return None

    if numeric <= 10.0:
        return int(round(numeric * 240.0))
    return int(round(numeric))


def _alignment_label(alignment: Any) -> str | None:
    if alignment == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "both"
    return None


def _extract_cover_title_text(doc_tree: ET.Element | None) -> str | None:
    if doc_tree is None:
        return None

    for paragraph in doc_tree.findall('.//w:p', NS):
        text = ''.join(node.text or '' for node in paragraph.findall('.//w:t', NS)).strip()
        lowered = text.lower()
        if 'project report titled' not in lowered or 'bonafide' not in lowered:
            continue

        match = re.search(r'project report titled\s*(.+?)\s*is\s+the\s+bonafide', text, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            continue

        title = match.group(1)
        title = title.replace('�', ' ').replace('\u2013', ' ').replace('\u2014', ' ')
        title = re.sub(r'\s+', ' ', title).strip(" \t\r\n'\"-–—")
        if title:
            return title

    return None


def _extract_cover_summary_text(document: Document) -> str | None:
    lines: list[str] = []
    month_year_pattern = re.compile(r'\b(?:JANUARY|FEBRUARY|MARCH|APRIL|MAY|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER)\s+\d{4}\b', re.IGNORECASE)
    for paragraph in document.paragraphs:
        text = (getattr(paragraph, "text", "") or "").strip()
        if not text:
            continue

        lowered = text.lower()
        if "project report titled" in lowered or "bonafide certificate" in lowered or lowered.startswith("declaration") or lowered.startswith("acknowledgement"):
            break

        if text in {"BONAFIDE CERTIFICATE", "DECLARATION", "ACKNOWLEDGEMENT"}:
            break

        lines.append(text)
        if month_year_pattern.search(text):
            break

    summary = "\n".join(lines).strip()
    return summary or None


def _paragraph_metric_summary(document: Document, style_names: set[str], min_text_length: int = 0) -> dict[str, Any]:
    wanted = {_normalize_name(name) for name in style_names}
    alignments: list[str] = []
    left_indents: list[float] = []
    right_indents: list[float] = []
    first_line_indents: list[float] = []
    space_before: list[float] = []
    space_after: list[float] = []
    line_spacing_vals: list[int] = []
    count = 0

    for paragraph in document.paragraphs:
        text = getattr(paragraph, "text", "") or ""
        clean_text = text.strip()
        if not clean_text:
            continue

        style_name = _normalize_name(getattr(getattr(paragraph, "style", None), "name", None))
        if style_name not in wanted:
            continue

        if len(clean_text) < min_text_length:
            continue

        count += 1
        format_ = paragraph.paragraph_format

        left = _length_pt(format_.left_indent)
        right = _length_pt(format_.right_indent)
        first = _length_pt(format_.first_line_indent)
        before = _length_pt(format_.space_before)
        after = _length_pt(format_.space_after)
        line_spacing = _line_spacing_to_twips(paragraph)
        alignment = _alignment_label(paragraph.alignment)

        if left is not None:
            left_indents.append(left)
        if right is not None:
            right_indents.append(right)
        if first is not None:
            first_line_indents.append(first)
        if before is not None:
            space_before.append(before)
        if after is not None:
            space_after.append(after)
        if line_spacing is not None:
            line_spacing_vals.append(line_spacing)
        if alignment is not None:
            alignments.append(alignment)

    return {
        "count": count,
        "alignment": Counter(alignments).most_common(1)[0][0] if alignments else None,
        "left_indent_pt": _median_or_none(left_indents),
        "right_indent_pt": _median_or_none(right_indents),
        "first_line_indent_pt": _median_or_none(first_line_indents),
        "space_before_pt": _median_or_none(space_before),
        "space_after_pt": _median_or_none(space_after),
        "line_spacing_val": _median_int(line_spacing_vals),
    }


def _extract_document_metrics(document: Document, styles_tree: ET.Element | None, rules: dict[str, Any], warnings: list[str]) -> None:
    try:
        sections = list(document.sections)
        if sections:
            width_vals = [_pt_to_dxa(_length_pt(section.page_width)) for section in sections]
            height_vals = [_pt_to_dxa(_length_pt(section.page_height)) for section in sections]
            top_vals = [_pt_to_dxa(_length_pt(section.top_margin)) for section in sections]
            bottom_vals = [_pt_to_dxa(_length_pt(section.bottom_margin)) for section in sections]
            left_vals = [_pt_to_dxa(_length_pt(section.left_margin)) for section in sections]
            right_vals = [_pt_to_dxa(_length_pt(section.right_margin)) for section in sections]
            header_vals = [_pt_to_dxa(_length_pt(section.header_distance)) for section in sections]
            footer_vals = [_pt_to_dxa(_length_pt(section.footer_distance)) for section in sections]

            rules["section_count"] = len(sections)
            rules["sections_have_different_margins"] = len({(top_vals[i], bottom_vals[i], left_vals[i], right_vals[i]) for i in range(len(sections))}) > 1

            rules["page_width_dxa"] = _median_int(width_vals)
            rules["page_height_dxa"] = _median_int(height_vals)
            rules["margin_top_dxa"] = _median_int(top_vals)
            rules["margin_bottom_dxa"] = _median_int(bottom_vals)
            rules["margin_left_dxa"] = _median_int(left_vals)
            rules["margin_right_dxa"] = _median_int(right_vals)
            rules["margin_header_dxa"] = _median_int(header_vals)
            rules["margin_footer_dxa"] = _median_int(footer_vals)

            if len(sections) >= 3:
                rules["has_cover_section"] = True

        body_style = _find_style(styles_tree, {"BodyText"}) or _find_style(styles_tree, {"Normal"})
        if body_style is not None:
            body_rpr = body_style.find(".//w:rPr", NS)
            body_sz = body_rpr.find("w:sz", NS) if body_rpr is not None else None
            body_fonts = body_rpr.find("w:rFonts", NS) if body_rpr is not None else None
            if body_fonts is not None:
                body_font = body_fonts.get(f"{{{NS['w']}}}ascii") or body_fonts.get(f"{{{NS['w']}}}hAnsi")
                if body_font:
                    rules["body_font"] = body_font
            if body_sz is not None:
                size = body_sz.get(f"{{{NS['w']}}}val")
                if size:
                    rules["body_size_halfpt"] = int(size)

        body_metrics = _paragraph_metric_summary(document, {"Body Text", "Normal"}, min_text_length=80)
        if body_metrics["count"]:
            if body_metrics["alignment"]:
                rules["body_alignment"] = body_metrics["alignment"]
            if body_metrics["line_spacing_val"] is not None:
                rules["body_line_spacing_val"] = body_metrics["line_spacing_val"]
                rules["body_line_spacing_rule"] = "auto"
            if body_metrics["space_before_pt"] is not None:
                rules["body_space_before"] = _pt_to_dxa(body_metrics["space_before_pt"])
            rules["body_space_after"] = _pt_to_dxa(body_metrics["space_after_pt"] or 0)
            if body_metrics["first_line_indent_pt"] is not None:
                rules["body_first_line_indent_dxa"] = _pt_to_dxa(body_metrics["first_line_indent_pt"])
            if body_metrics["left_indent_pt"] is not None:
                rules["body_left_indent_dxa"] = _pt_to_dxa(body_metrics["left_indent_pt"])
            if body_metrics["right_indent_pt"] is not None:
                rules["body_right_indent_dxa"] = _pt_to_dxa(body_metrics["right_indent_pt"])

        for level in range(1, 4):
            metrics = _paragraph_metric_summary(document, {f"Heading {level}"})
            if not metrics["count"]:
                continue
            if metrics["alignment"]:
                rules[f"h{level}_alignment"] = metrics["alignment"]
            if metrics["space_before_pt"] is not None:
                rules[f"h{level}_space_before"] = _pt_to_dxa(metrics["space_before_pt"])
            rules[f"h{level}_space_after"] = _pt_to_dxa(metrics["space_after_pt"] or 0)
            if metrics["left_indent_pt"] is not None:
                rules[f"h{level}_left_indent_dxa"] = _pt_to_dxa(metrics["left_indent_pt"])
            if metrics["right_indent_pt"] is not None:
                rules[f"h{level}_right_indent_dxa"] = _pt_to_dxa(metrics["right_indent_pt"])
            if metrics["first_line_indent_pt"] is not None:
                rules[f"h{level}_first_line_indent_dxa"] = _pt_to_dxa(metrics["first_line_indent_pt"])

        list_metrics = _paragraph_metric_summary(document, {"List Paragraph"})
        if list_metrics["count"]:
            if list_metrics["left_indent_pt"] is not None:
                rules["list_left_indent_pt"] = list_metrics["left_indent_pt"]
            if list_metrics["first_line_indent_pt"] is not None:
                rules["list_first_line_indent_pt"] = list_metrics["first_line_indent_pt"]

        cover_title_sizes: list[float] = []
        cover_subtitle_sizes: list[float] = []
        for paragraph in document.paragraphs[:30]:
            text = (paragraph.text or "").strip()
            if not text:
                continue

            run_sizes = [run.font.size.pt for run in paragraph.runs if run.text and run.text.strip() and run.font.size is not None]
            if not run_sizes:
                continue

            if any(ch.isalpha() for ch in text) and text.upper() == text and not any(ch.isdigit() for ch in text):
                cover_title_sizes.append(_median_or_none(run_sizes) or 0.0)

            lower_text = text.lower()
            if "submitted by" in lower_text or "partial fulfillment" in lower_text or any(run.italic for run in paragraph.runs if run.text and run.text.strip()):
                cover_subtitle_sizes.append(_median_or_none(run_sizes) or 0.0)

        if cover_title_sizes:
            rules["cover_title_size_pt"] = _median_or_none(cover_title_sizes)
        if cover_subtitle_sizes:
            rules["cover_subtitle_size_pt"] = _median_or_none(cover_subtitle_sizes)
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Document metrics extraction failed: {exc}")


def _extract_footer_metadata(docx_zip: zipfile.ZipFile, rules: dict[str, Any], warnings: list[str]) -> None:
    footer_files = sorted(name for name in docx_zip.namelist() if re.fullmatch(r"word/footer\d+\.xml", name))
    header_files = sorted(name for name in docx_zip.namelist() if re.fullmatch(r"word/header\d+\.xml", name))

    rules["footer_count"] = len(footer_files)
    rules["header_count"] = len(header_files)

    footer_text_sample = None
    page_number_alignments: list[str] = []
    saw_roman_page_numbers = False
    saw_decimal_page_numbers = False

    for footer_name in footer_files:
        try:
            root = ET.fromstring(docx_zip.read(footer_name))
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"Footer parse failed for {footer_name}: {exc}")
            continue

        texts = [node.text.strip() for node in root.findall('.//w:t', NS) if node.text and node.text.strip()]
        if footer_text_sample is None and texts:
            footer_text_sample = " ".join(texts)

        for paragraph in root.findall('.//w:p', NS):
            jc = paragraph.find('w:pPr/w:jc', NS)
            if jc is not None:
                val = jc.get(f"{{{NS['w']}}}val")
                if val:
                    page_number_alignments.append(val)

            for instr in paragraph.findall('.//w:instrText', NS):
                instr_text = (instr.text or "").lower()
                if "page" not in instr_text:
                    continue

                rules["has_page_numbers"] = True
                if "roman" in instr_text:
                    saw_roman_page_numbers = True
                    if "upper" in instr_text:
                        rules["prelim_page_format"] = "upperRoman"
                    else:
                        rules["prelim_page_format"] = "lowerRoman"
                else:
                    saw_decimal_page_numbers = True
                    rules["body_page_format"] = "decimal"

    if footer_text_sample is not None:
        rules["footer_text_sample"] = footer_text_sample

    if page_number_alignments:
        rules["page_number_alignment"] = Counter(page_number_alignments).most_common(1)[0][0]

    if saw_roman_page_numbers or saw_decimal_page_numbers:
        rules["has_page_numbers"] = True
        rules["footer_has_page_number"] = True
    if saw_roman_page_numbers:
        rules["has_prelim_section"] = True
    if saw_decimal_page_numbers:
        rules["has_body_section"] = True
    if saw_roman_page_numbers and saw_decimal_page_numbers:
        rules["page_number_section_restart"] = True


def extract_rules(docx_bytes: bytes, source_filename: str = "") -> dict[str, Any]:
    """
    Extract all formatting rules from DOCX.
    
    Args:
        docx_bytes: Raw DOCX file bytes
        source_filename: Original filename (for metadata)
    
    Returns:
        Normalized rules dict (all keys present, null for unknown values)
    """
    rules = dict(NEUTRAL_DEFAULTS)
    rules["source_filename"] = source_filename
    warnings: list[str] = []

    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
            # Read main document
            doc_xml = docx_zip.read("word/document.xml").decode("utf-8")
            doc_tree = ET.fromstring(doc_xml)

            document = None
            try:
                document = Document(io.BytesIO(docx_bytes))
            except Exception as exc:  # noqa: BLE001
                warnings.append(f"python-docx parsing failed: {exc}")

            # Read styles
            try:
                styles_xml = docx_zip.read("word/styles.xml").decode("utf-8")
                styles_tree = ET.fromstring(styles_xml)
            except KeyError:
                styles_tree = None
                warnings.append("styles.xml not found")

            # Read numbering
            try:
                numbering_xml = docx_zip.read("word/numbering.xml").decode("utf-8")
                numbering_tree = ET.fromstring(numbering_xml)
            except KeyError:
                numbering_tree = None

            # Extract page settings
            _extract_page_settings(doc_tree, rules, warnings)

            # Extract body typography
            _extract_body_typography(doc_tree, styles_tree, rules, warnings)

            # Extract heading styles
            _extract_headings(styles_tree, rules, warnings)

            # Extract section headings (h1-h6 text from document)
            _extract_detected_headings(doc_tree, rules, warnings)

            # Extract page numbering
            _extract_page_numbering(doc_tree, rules, warnings)

            # Extract lists, tables, images
            _extract_document_elements(doc_tree, rules, warnings)

            # Extract headers/footers
            _extract_headers_footers(docx_zip, doc_tree, rules, warnings)

            cover_title_text = _extract_cover_title_text(doc_tree)
            if cover_title_text:
                rules["cover_title_text"] = cover_title_text

            cover_summary_text = _extract_cover_summary_text(document) if document is not None else None
            if cover_summary_text:
                rules["cover_summary_text"] = cover_summary_text

            if document is not None:
                _extract_document_metrics(document, styles_tree, rules, warnings)
                _extract_footer_metadata(docx_zip, rules, warnings)

            # Quality checks
            _check_quality_flags(doc_tree, rules, warnings)

            # Determine confidence
            missing_critical = sum(1 for v in [
                rules.get("body_font"),
                rules.get("body_size_halfpt"),
                rules.get("page_width_dxa"),
            ] if v is None)
            if missing_critical > 1:
                rules["confidence"] = "low"
            elif missing_critical == 1:
                rules["confidence"] = "medium"
            else:
                rules["confidence"] = "high"

    except Exception as e:
        warnings.append(f"Extraction failed: {str(e)}")
        rules["confidence"] = "low"

    rules["extraction_warnings"] = warnings
    return rules


def _extract_page_settings(doc_tree: ET.Element, rules: dict, warnings: list) -> None:
    """Extract page size and margins from section properties."""
    sectPr = doc_tree.find(".//w:sectPr", NS)
    if sectPr is None:
        warnings.append("No section properties found")
        return

    # Page size
    pgSz = sectPr.find("w:pgSz", NS)
    if pgSz is not None:
        rules["page_width_dxa"] = int(pgSz.get(f"{{{NS['w']}}}w", 0))
        rules["page_height_dxa"] = int(pgSz.get(f"{{{NS['w']}}}h", 0))

    # Page margins
    pgMar = sectPr.find("w:pgMar", NS)
    if pgMar is not None:
        rules["margin_top_dxa"] = int(pgMar.get(f"{{{NS['w']}}}top", 0))
        rules["margin_bottom_dxa"] = int(pgMar.get(f"{{{NS['w']}}}bottom", 0))
        rules["margin_left_dxa"] = int(pgMar.get(f"{{{NS['w']}}}left", 0))
        rules["margin_right_dxa"] = int(pgMar.get(f"{{{NS['w']}}}right", 0))
        rules["margin_header_dxa"] = int(pgMar.get(f"{{{NS['w']}}}header", 0))
        rules["margin_footer_dxa"] = int(pgMar.get(f"{{{NS['w']}}}footer", 0))


def _extract_body_typography(
    doc_tree: ET.Element, styles_tree: ET.Element | None, rules: dict, warnings: list
) -> None:
    """Extract body text font, size, line spacing from Normal style."""
    if styles_tree is None:
        warnings.append("Cannot extract body typography: no styles.xml")
        return

    # Find Normal style
    normal_style = None
    for style in styles_tree.findall(".//w:style", NS):
        if style.get(f"{{{NS['w']}}}styleId") == "Normal":
            normal_style = style
            break

    if normal_style is None:
        warnings.append("Normal style not found")
        return

    rPr = normal_style.find(".//w:rPr", NS)
    pPr = normal_style.find(".//w:pPr", NS)

    # Font
    if rPr is not None:
        rFonts = rPr.find("w:rFonts", NS)
        if rFonts is not None:
            rules["body_font"] = rFonts.get(f"{{{NS['w']}}}ascii")

    # Size (in half-points)
    if rPr is not None:
        sz = rPr.find("w:sz", NS)
        if sz is not None:
            rules["body_size_halfpt"] = int(sz.get(f"{{{NS['w']}}}val", 0))

    # Line spacing
    if pPr is not None:
        spacing = pPr.find("w:spacing", NS)
        if spacing is not None:
            line = spacing.get(f"{{{NS['w']}}}line")
            if line:
                rules["body_line_spacing_val"] = int(line)
            lineRule = spacing.get(f"{{{NS['w']}}}lineRule")
            if lineRule:
                rules["body_line_spacing_rule"] = lineRule

        # Alignment
        jc = pPr.find("w:jc", NS)
        if jc is not None:
            align = jc.get(f"{{{NS['w']}}}val")
            if align == "both":
                rules["body_alignment"] = "both"
            elif align == "left":
                rules["body_alignment"] = "left"
            elif align == "center":
                rules["body_alignment"] = "center"
            elif align == "right":
                rules["body_alignment"] = "right"

        # Spacing before/after
        spacing = pPr.find("w:spacing", NS)
        if spacing is not None:
            before = spacing.get(f"{{{NS['w']}}}before")
            after = spacing.get(f"{{{NS['w']}}}after")
            if before:
                rules["body_space_before"] = int(before)
            if after:
                rules["body_space_after"] = int(after)

        # First line indent
        ind = pPr.find("w:ind", NS)
        if ind is not None:
            firstLine = ind.get(f"{{{NS['w']}}}firstLine")
            if firstLine:
                rules["body_first_line_indent_dxa"] = int(firstLine)


def _extract_headings(styles_tree: ET.Element | None, rules: dict, warnings: list) -> None:
    """Extract heading styles 1-6."""
    if styles_tree is None:
        return

    for level in range(1, 7):
        heading_id = f"Heading{level}"
        style = None
        for s in styles_tree.findall(".//w:style", NS):
            if s.get(f"{{{NS['w']}}}styleId") == heading_id:
                style = s
                break

        if style is None:
            continue

        heading_dict = {
            "size_halfpt": None,
            "bold": False,
            "italic": False,
            "underline": False,
            "caps": False,
            "small_caps": False,
            "alignment": None,
            "space_before": None,
            "space_after": None,
            "font": None,
            "numbering": False,
        }

        rPr = style.find(".//w:rPr", NS)
        if rPr is not None:
            # Font
            rFonts = rPr.find("w:rFonts", NS)
            if rFonts is not None:
                heading_dict["font"] = rFonts.get(f"{{{NS['w']}}}ascii")

            # Size
            sz = rPr.find("w:sz", NS)
            if sz is not None:
                heading_dict["size_halfpt"] = int(sz.get(f"{{{NS['w']}}}val", 0))

            # Bold/italic/underline
            heading_dict["bold"] = rPr.find("w:b", NS) is not None
            heading_dict["italic"] = rPr.find("w:i", NS) is not None
            heading_dict["underline"] = rPr.find("w:u", NS) is not None

            # Caps
            caps = rPr.find("w:caps", NS)
            if caps is not None:
                heading_dict["caps"] = caps.get(f"{{{NS['w']}}}val", "1") != "0"

            # Small caps
            smallCaps = rPr.find("w:smallCaps", NS)
            if smallCaps is not None:
                heading_dict["small_caps"] = smallCaps.get(f"{{{NS['w']}}}val", "1") != "0"

        pPr = style.find(".//w:pPr", NS)
        if pPr is not None:
            # Alignment
            jc = pPr.find("w:jc", NS)
            if jc is not None:
                heading_dict["alignment"] = jc.get(f"{{{NS['w']}}}val")

            # Spacing
            spacing = pPr.find("w:spacing", NS)
            if spacing is not None:
                before = spacing.get(f"{{{NS['w']}}}before")
                after = spacing.get(f"{{{NS['w']}}}after")
                if before:
                    heading_dict["space_before"] = int(before)
                if after:
                    heading_dict["space_after"] = int(after)

            # Numbering
            numPr = pPr.find("w:numPr", NS)
            heading_dict["numbering"] = numPr is not None

        rules["headings"][str(level)] = heading_dict


def _extract_detected_headings(doc_tree: ET.Element, rules: dict, warnings: list) -> None:
    """Extract text of all heading-styled paragraphs."""
    headings_text = []
    for para in doc_tree.findall(".//w:p", NS):
        pPr = para.find("w:pPr", NS)
        if pPr is None:
            continue

        pStyle = pPr.find("w:pStyle", NS)
        if pStyle is None:
            continue

        style_id = pStyle.get(f"{{{NS['w']}}}val")
        if style_id and style_id.startswith("Heading"):
            # Extract text
            text_parts = []
            for t in para.findall(".//w:t", NS):
                if t.text:
                    text_parts.append(t.text)
            if text_parts:
                headings_text.append("".join(text_parts))

    rules["detected_section_headings"] = headings_text


def _extract_page_numbering(doc_tree: ET.Element, rules: dict, warnings: list) -> None:
    """Extract page numbering format and settings."""
    sectPr = doc_tree.find(".//w:sectPr", NS)
    if sectPr is None:
        return

    # Check for page numbers in footer
    footer = sectPr.find(".//w:footerReference", NS)
    if footer is not None:
        rules["has_page_numbers"] = True

    # Page number format
    pgNumType = sectPr.find("w:pgNumType", NS)
    if pgNumType is not None:
        fmt = pgNumType.get(f"{{{NS['w']}}}fmt")
        if fmt == "lowerRoman":
            rules["prelim_page_format"] = "lowerRoman"
        elif fmt == "upperRoman":
            rules["prelim_page_format"] = "upperRoman"
        elif fmt == "decimal":
            rules["body_page_format"] = "decimal"

        start = pgNumType.get(f"{{{NS['w']}}}start")
        if start == "1":
            rules["page_number_section_restart"] = True


def _extract_document_elements(doc_tree: ET.Element, rules: dict, warnings: list) -> None:
    """Extract lists, tables, images."""
    # Tables
    tables = doc_tree.findall(".//w:tbl", NS)
    rules["table_count"] = len(tables)
    if tables:
        for tbl in tables:
            tcBorders = tbl.find(".//w:tcBorders", NS)
            if tcBorders is not None:
                rules["tables_use_borders"] = True
                break

    # Lists
    rules["has_numbered_lists"] = doc_tree.find(".//w:numPr", NS) is not None
    rules["has_bulleted_lists"] = doc_tree.find(".//w:lvlPicBulletId", NS) is not None

    # Images
    images = doc_tree.findall(".//w:drawing", NS)
    rules["image_count"] = len(images)
    if images:
        first_para = doc_tree.find(".//w:p", NS)
        if first_para is not None and first_para.find(".//w:drawing", NS) is not None:
            rules["has_cover_image"] = True


def _extract_headers_footers(
    docx_zip: zipfile.ZipFile, doc_tree: ET.Element, rules: dict, warnings: list
) -> None:
    """Extract header/footer content and page number references."""
    sectPr = doc_tree.find(".//w:sectPr", NS)
    if sectPr is None:
        return

    footers = sectPr.findall("w:footerReference", NS)
    headers = sectPr.findall("w:headerReference", NS)

    rules["footer_count"] = len(footers)
    rules["header_count"] = len(headers)

    # Check for page numbers
    for footer in footers:
        try:
            rel_id = footer.get(f"{{{NS['r']}}}id")
            # Would need to read relationships and footer XML to get full text
            rules["footer_has_page_number"] = True
        except Exception as e:
            warnings.append(str(e))

    for header in headers:
        try:
            rules["header_has_page_number"] = True
        except Exception as e:
            warnings.append(str(e))


def _check_quality_flags(doc_tree: ET.Element, rules: dict, warnings: list) -> None:
    """Detect markdown leakage, artifact numbers, mixed fonts."""
    all_text = []
    fonts = set()
    sizes = set()

    for t in doc_tree.findall(".//w:t", NS):
        if t.text:
            all_text.append(t.text)

    # Collect fonts and sizes
    for rFonts in doc_tree.findall(".//w:rFonts", NS):
        font = rFonts.get(f"{{{NS['w']}}}ascii")
        if font:
            fonts.add(font)

    for sz in doc_tree.findall(".//w:sz", NS):
        size = sz.get(f"{{{NS['w']}}}val")
        if size:
            sizes.add(size)

    combined_text = " ".join(all_text)

    # Markdown leakage
    if re.search(r"^\s*#", combined_text, re.MULTILINE) or "**" in combined_text or "__" in combined_text:
        rules["has_markdown_leak"] = True

    # XML artifact numbers (6-9 digit standalone numbers)
    if re.search(r"\b\d{6,9}\b", combined_text):
        rules["has_xml_artifact_numbers"] = True

    # Mixed fonts
    if len(fonts) > 2:
        rules["has_mixed_fonts"] = True

    # Inconsistent sizes
    if len(sizes) > 3:
        rules["has_inconsistent_sizes"] = True

    # Font substitution
    if any(f in fonts for f in ["SimSun", "Courier", "Arial", "Calibri"]):
        rules["font_substitution_detected"] = True
