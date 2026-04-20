from __future__ import annotations

import os
import zipfile
from dataclasses import dataclass
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

from .constants import FIXED_EPOCH
from .docx_fonts import apply_run_font_override
from .models import DocumentSpec, ParagraphElement, TableElement
from .template_registry import TemplateVersion
from .units import pt_to_emu, to_emu


@dataclass(frozen=True)
class AtomicWriteResult:
    output_path: Path
    bytes_written: int


def _set_fixed_table_layout(table) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_cell_width(cell, width_emu: int) -> None:
    cell.width = Emu(width_emu)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    twips = int(round(width_emu / 635))
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(twips))


def _style_for(element: dict[str, Any], resolved_styles: dict[str, dict[str, Any]]) -> dict[str, Any]:
    style_name = str(element.get("style") or "Normal")
    style = resolved_styles.get(style_name) or resolved_styles.get("Normal")
    if style is None:
        raise ValueError(f"Missing resolved style: {style_name}")
    return style


def _apply_page_layout(doc: Document, template: TemplateVersion) -> None:
    section = doc.sections[0]
    layout = template.page_layout
    margins = layout.get("margins", {})

    width_in = layout["size"]["width_in"]
    height_in = layout["size"]["height_in"]

    section.page_width = Emu(to_emu(width_in, "in"))
    section.page_height = Emu(to_emu(height_in, "in"))
    section.top_margin = Emu(to_emu(margins.get("top_in", 1.0), "in"))
    section.right_margin = Emu(to_emu(margins.get("right_in", 1.0), "in"))
    section.bottom_margin = Emu(to_emu(margins.get("bottom_in", 1.0), "in"))
    section.left_margin = Emu(to_emu(margins.get("left_in", 1.0), "in"))


def _apply_header_footer(doc: Document, template: TemplateVersion) -> None:
    section = doc.sections[0]
    hf = template.header_footer
    header_text = str((hf.get("header") or {}).get("text") or "")
    footer_text = str((hf.get("footer") or {}).get("text") or "")

    section.header.paragraphs[0].text = header_text
    section.footer.paragraphs[0].text = footer_text


def _set_fixed_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.author = "DocuForageAI"
    props.last_modified_by = "DocuForageAI"
    props.created = FIXED_EPOCH
    props.modified = FIXED_EPOCH
    props.revision = 1
    props.title = "Generated Document"


def _add_paragraph(doc: Document, element: dict[str, Any], style: dict[str, Any]) -> None:
    paragraph_model = ParagraphElement.model_validate(element)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.page_break_before = bool(element.get("break_before", False))
    paragraph.paragraph_format.line_spacing = float(style.get("line_spacing", 1.0))
    paragraph.paragraph_format.space_before = Pt(float(style.get("space_before_pt", 0.0)))
    paragraph.paragraph_format.space_after = Pt(float(style.get("space_after_pt", 0.0)))

    run = paragraph.add_run(paragraph_model.text)
    font_name = str(style["font_name"])
    run.font.size = Pt(float(style["font_size_pt"]))
    apply_run_font_override(run, font_name)


def _add_table(doc: Document, element: dict[str, Any], break_before: bool) -> None:
    table_model = TableElement.model_validate(element)
    if break_before:
        br = doc.add_paragraph()
        br.paragraph_format.page_break_before = True

    col_count = len(table_model.rows[0].cells)
    table = doc.add_table(rows=len(table_model.rows), cols=col_count)
    _set_fixed_table_layout(table)

    for r_idx, row in enumerate(table_model.rows):
        for c_idx, cell_spec in enumerate(row.cells):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_spec.text
            width_emu = cell_spec.width.as_emu()
            _set_cell_width(cell, width_emu)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _build_document_from_annotated(
    spec: DocumentSpec,
    template: TemplateVersion,
    annotated_elements: list[dict[str, Any]],
    resolved_styles: dict[str, dict[str, Any]],
) -> Document:
    doc = Document()
    _apply_page_layout(doc, template)
    _apply_header_footer(doc, template)

    for element in annotated_elements:
        style = _style_for(element, resolved_styles)
        element_type = str(element.get("type"))
        if element_type == "paragraph":
            _add_paragraph(doc, element, style)
        elif element_type == "table":
            _add_table(doc, element, break_before=bool(element.get("break_before", False)))
        else:
            # Images are phase-6+ concern in this pipeline version. Keep deterministic text marker.
            marker = doc.add_paragraph()
            run = marker.add_run(f"[unsupported element type: {element_type}]")
            run.font.size = Pt(float(style.get("font_size_pt", 11)))
            apply_run_font_override(run, str(style.get("font_name", "Calibri")))

    _set_fixed_core_properties(doc)
    return doc


def _canonicalize_docx(path: Path) -> None:
    fixed_dt = (2000, 1, 1, 0, 0, 0)
    with zipfile.ZipFile(path, "r") as src:
        entries = [(info.filename, src.read(info.filename)) for info in src.infolist()]

    entries.sort(key=lambda item: item[0])

    canonical = path.with_suffix(path.suffix + ".canonical")
    with zipfile.ZipFile(canonical, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as dst:
        for filename, data in entries:
            info = zipfile.ZipInfo(filename)
            info.date_time = fixed_dt
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0
            dst.writestr(info, data)

    os.replace(canonical, path)


def write_docx_atomic(
    *,
    spec: DocumentSpec,
    template: TemplateVersion,
    annotated_elements: list[dict[str, Any]],
    resolved_styles: dict[str, dict[str, Any]],
    output_path: Path,
) -> AtomicWriteResult:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with NamedTemporaryFile(prefix=output_path.stem + ".", suffix=".tmp.docx", dir=str(output_path.parent), delete=False) as fh:
        tmp_path = Path(fh.name)

    try:
        doc = _build_document_from_annotated(spec, template, annotated_elements, resolved_styles)
        doc.save(tmp_path)
        _canonicalize_docx(tmp_path)

        bytes_written = tmp_path.stat().st_size
        os.replace(tmp_path, output_path)
        return AtomicWriteResult(output_path=output_path, bytes_written=bytes_written)
    except Exception:
        if tmp_path.exists():
            tmp_path.unlink()
        raise
